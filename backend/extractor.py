import io
from pathlib import Path
import re


def extract_text_from_file(content: bytes, filename: str) -> str:
    ext = Path(filename).suffix.lower()

    if ext == ".pdf":
        return _extract_pdf(content)
    elif ext in (".doc", ".docx"):
        return _extract_word(content)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Please upload PDF or Word (.doc/.docx).")


def _extract_pdf(content: bytes) -> str:
    try:
        import fitz
        with fitz.open(stream=content, filetype="pdf") as doc:
            text_parts = []
            for page in doc:
                text_parts.append(page.get_text("text") or "")


        text = "\n".join(text_parts).strip()
        text = re.sub(r'\n+', '\n', text)

        if not text:
            raise ValueError("PDF appears to be scanned/image-only. No text could be extracted.")
        return text
    except ImportError:
        return _extract_pdf_fallback(content)


def _extract_pdf_fallback(content: bytes) -> str:

    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
        text = "\n".join(pages).strip()
        if not text:
            raise ValueError("Could not extract text from PDF.")
        return text
    except ImportError:
        raise ImportError(
            "No PDF library found. Install one:\n"
            "  pip install PyMuPDF\n"
            "  OR\n"
            "  pip install pdfplumber"
        )


def _extract_word(content: bytes) -> str:

    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        text = "\n".join(paragraphs).strip()
        text = re.sub(r'\n+', '\n', text)
        if not text:
            raise ValueError("Word file appears to be empty or has no readable text.")
        return text
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")
